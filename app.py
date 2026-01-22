"""
Streamlit OCR Application
Extracts text from images using OCR and allows downloading as Word document
"""

import streamlit as st
import easyocr
from PIL import Image
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime

# Page configuration
st.set_page_config(
    page_title="OCR Text Extractor",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
        padding: 0.5rem;
        font-size: 16px;
        border-radius: 8px;
        border: none;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        background-color: #45a049;
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }
    .upload-section {
        background-color: #f0f2f6;
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .extracted-text {
        background-color: #ffffff;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #4CAF50;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    h1 {
        color: #1f77b4;
        text-align: center;
        padding-bottom: 1rem;
    }
    .success-message {
        background-color: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'extracted_texts' not in st.session_state:
    st.session_state.extracted_texts = []
if 'image_names' not in st.session_state:
    st.session_state.image_names = []

@st.cache_resource
def load_ocr_reader():
    """Load EasyOCR reader (cached to avoid reloading)"""
    try:
        reader = easyocr.Reader(['en'], gpu=False)
        return reader
    except Exception as e:
        st.error(f"Error loading OCR reader: {str(e)}")
        return None

def extract_text_from_image(image, reader):
    """Extract text from image using EasyOCR"""
    try:
        # Convert PIL Image to numpy array
        image_np = np.array(image)
        
        # Perform OCR
        results = reader.readtext(image_np)
        
        # Extract text from results
        extracted_text = '\n'.join([text[1] for text in results])
        
        return extracted_text
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return ""

def create_word_document(texts, image_names):
    """Create a Word document with extracted text"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('OCR Extracted Text', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.runs[0]
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Add spacing
    
    # Add extracted text for each image
    for idx, (text, img_name) in enumerate(zip(texts, image_names), 1):
        # Add image heading
        heading = doc.add_heading(f'Image {idx}: {img_name}', level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(31, 119, 180)
        
        # Add extracted text
        if text.strip():
            paragraph = doc.add_paragraph(text)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5
            
            # Style the text
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        else:
            no_text = doc.add_paragraph("No text extracted from this image.")
            no_text_run = no_text.runs[0]
            no_text_run.font.italic = True
            no_text_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Add separator
        if idx < len(texts):
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
    
    # Save to BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# Main App
def main():
    # Title and description
    st.title("📝 OCR Text Extractor")
    st.markdown("""
        <p style='text-align: center; font-size: 18px; color: #666;'>
            Upload images to extract text using OCR technology and download as Word document
        </p>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("📋 Instructions")
        st.markdown("""
        1. **Upload Images**: Click 'Browse files' to upload one or more images
        2. **Extract Text**: Click 'Extract Text from Images' button
        3. **Review**: Check the extracted text displayed below
        4. **Download**: Click 'Download as Word Document' to save
        
        ---
        
        ### 📌 Supported Formats
        - JPG/JPEG
        - PNG
        - BMP
        
        ### 💡 Tips
        - Use clear, high-resolution images
        - Ensure good lighting and contrast
        - Avoid blurry or skewed images
        """)
        
        st.info("🚀 Powered by EasyOCR & Streamlit")
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("<div class='upload-section'>", unsafe_allow_html=True)
        st.subheader("📤 Upload Images")
        
        uploaded_files = st.file_uploader(
            "Choose image files",
            type=['png', 'jpg', 'jpeg', 'bmp'],
            accept_multiple_files=True,
            help="Upload one or more images for text extraction"
        )
        st.markdown("</div>", unsafe_allow_html=True)
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} image(s) uploaded successfully!")
            
            # Display uploaded images
            st.subheader("🖼️ Uploaded Images")
            for idx, uploaded_file in enumerate(uploaded_files, 1):
                with st.expander(f"Image {idx}: {uploaded_file.name}"):
                    image = Image.open(uploaded_file)
                    st.image(image, use_column_width=True)
    
    with col2:
        if uploaded_files:
            st.markdown("<div class='upload-section'>", unsafe_allow_html=True)
            st.subheader("⚙️ OCR Processing")
            
            if st.button("🔍 Extract Text from Images", type="primary"):
                # Load OCR reader
                reader = load_ocr_reader()
                
                if reader is None:
                    st.error("Failed to load OCR reader. Please try again.")
                    return
                
                # Clear previous results
                st.session_state.extracted_texts = []
                st.session_state.image_names = []
                
                # Progress bar
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Process each image
                for idx, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"Processing {uploaded_file.name}...")
                    
                    # Open image
                    image = Image.open(uploaded_file)
                    
                    # Extract text
                    extracted_text = extract_text_from_image(image, reader)
                    
                    # Store results
                    st.session_state.extracted_texts.append(extracted_text)
                    st.session_state.image_names.append(uploaded_file.name)
                    
                    # Update progress
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                
                status_text.empty()
                progress_bar.empty()
                
                st.markdown("<div class='success-message'>", unsafe_allow_html=True)
                st.write("✨ Text extraction completed successfully!")
                st.markdown("</div>", unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
    
    # Display extracted text
    if st.session_state.extracted_texts:
        st.markdown("---")
        st.header("📄 Extracted Text")
        
        # Create tabs for each image
        if len(st.session_state.extracted_texts) > 1:
            tabs = st.tabs([f"Image {i+1}" for i in range(len(st.session_state.extracted_texts))])
            
            for idx, tab in enumerate(tabs):
                with tab:
                    st.markdown("<div class='extracted-text'>", unsafe_allow_html=True)
                    st.subheader(f"📝 {st.session_state.image_names[idx]}")
                    
                    text = st.session_state.extracted_texts[idx]
                    if text.strip():
                        st.text_area(
                            "Extracted Text",
                            value=text,
                            height=300,
                            key=f"text_{idx}",
                            label_visibility="collapsed"
                        )
                        st.caption(f"📊 Characters: {len(text)} | Words: {len(text.split())}")
                    else:
                        st.warning("No text was extracted from this image.")
                    st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.markdown("<div class='extracted-text'>", unsafe_allow_html=True)
            st.subheader(f"📝 {st.session_state.image_names[0]}")
            
            text = st.session_state.extracted_texts[0]
            if text.strip():
                st.text_area(
                    "Extracted Text",
                    value=text,
                    height=300,
                    label_visibility="collapsed"
                )
                st.caption(f"📊 Characters: {len(text)} | Words: {len(text.split())}")
            else:
                st.warning("No text was extracted from this image.")
            st.markdown("</div>", unsafe_allow_html=True)
        
        # Download section
        st.markdown("---")
        st.header("💾 Download Options")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            # Create Word document
            doc_io = create_word_document(
                st.session_state.extracted_texts,
                st.session_state.image_names
            )
            
            # Download button
            st.download_button(
                label="📥 Download as Word Document",
                data=doc_io,
                file_name=f"OCR_Extracted_Text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download all extracted text as a formatted Word document"
            )

if __name__ == "__main__":
    main()
